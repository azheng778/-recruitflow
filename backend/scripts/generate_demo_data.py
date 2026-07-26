from __future__ import annotations

import hashlib
import random
import uuid
from collections import Counter
from datetime import datetime, timedelta
from decimal import Decimal
from io import BytesIO
from pathlib import Path

from docx import Document
from sqlalchemy import func, select, text

from app.config import settings
from app.database import SessionLocal, engine
from app.models import (
    AgentConversation,
    AgentMessage,
    ApprovalRequest,
    AuditLog,
    Candidate,
    CandidateJob,
    CandidateStatusHistory,
    Communication,
    DocumentSyncJob,
    InboundEvent,
    Interview,
    Job,
    Notification,
    Resume,
    User,
    utcnow,
)
from app.security import hash_password


SEED = 20260726
RANDOM = random.Random(SEED)
NAMESPACE = uuid.UUID("1bea42fc-f857-4c31-8604-215020f181fe")

TARGETS = {
    "users": 12,
    "jobs": 10,
    "candidates": 150,
    "candidate_jobs": 180,
    "resumes": 110,
    "interviews": 100,
    "communications": 500,
    "notifications": 160,
    "audit_logs": 1200,
    "agent_conversations": 40,
    "agent_messages": 240,
    "approval_requests": 35,
    "inbound_events": 250,
    "document_sync_jobs": 500,
}

STAGE_TARGETS = {
    "new": 27,
    "screening": 36,
    "interview_1": 31,
    "interview_2": 22,
    "final_interview": 14,
    "on_hold": 9,
    "offer": 11,
    "hired": 9,
    "rejected": 16,
    "withdrawn": 5,
}

SURNAMES = "赵钱孙李周吴郑王冯陈褚卫蒋沈韩杨朱秦尤许何吕施张孔曹严华金魏陶姜"
GIVEN_NAMES = [
    "安然", "博文", "晨曦", "德宇", "恩泽", "飞扬", "嘉宁", "浩然", "静怡", "凯文",
    "乐天", "明远", "宁馨", "启航", "若溪", "思源", "天宇", "文清", "欣妍", "逸凡",
    "雨桐", "子墨", "昊宇", "婉清", "景行", "可欣", "凌云", "梦琪", "清越", "书航",
]
SKILL_GROUPS = [
    ["Python", "FastAPI", "MySQL"],
    ["LangChain", "LangGraph", "Python"],
    ["Vue", "JavaScript", "CSS"],
    ["React", "TypeScript", "ECharts"],
    ["Java", "Spring Boot", "SQL"],
    ["Docker", "Kubernetes", "Linux"],
    ["数据分析", "SQL", "Excel"],
    ["招聘运营", "企业微信", "沟通协调"],
]


def sid(kind: str, value: object) -> str:
    return str(uuid.uuid5(NAMESPACE, f"recruitflow-demo:{kind}:{value}"))


def digest(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def count_rows(db, model) -> int:
    return int(db.scalar(select(func.count()).select_from(model)) or 0)


def stage_chain(stage: str, index: int) -> list[str]:
    base = ["new", "screening", "interview_1", "interview_2", "final_interview"]
    if stage == "new":
        return ["new"]
    if stage == "screening":
        return base[:2]
    if stage == "interview_1":
        return base[:3]
    if stage == "interview_2":
        return base[:4]
    if stage == "final_interview":
        return base
    if stage == "on_hold":
        return ["new", "screening", "on_hold"]
    if stage == "offer":
        return base + ["offer"]
    if stage == "hired":
        return base + ["offer", "hired"]
    if stage in {"rejected", "withdrawn"}:
        terminal_at = 1 + index % 4
        return base[: terminal_at + 1] + [stage]
    raise ValueError(f"Unsupported stage: {stage}")


def next_action(stage: str) -> str | None:
    return {
        "new": "完成简历初筛",
        "screening": "确认一面安排",
        "interview_1": "收集一面反馈",
        "interview_2": "准备终面",
        "final_interview": "发起录用决策",
        "on_hold": "七天后重新评估",
        "offer": "跟进录用意向",
    }.get(stage)


def ensure_users(db) -> None:
    missing = max(0, TARGETS["users"] - count_rows(db, User))
    if not missing:
        return
    password_hash = hash_password("RecruitFlow!2026")
    for index in range(1, 50):
        if missing <= 0:
            break
        user_id = sid("user", index)
        if db.get(User, user_id):
            continue
        role = "hr" if index <= 3 else "interviewer"
        db.add(
            User(
                id=user_id,
                username=f"demo_{role}_{index:02d}",
                display_name=f"演示{'招聘专员' if role == 'hr' else '面试官'}{index:02d}",
                email=f"demo-user-{index:02d}@example.com",
                phone=f"1379000{index:04d}",
                password_hash=password_hash,
                role=role,
                department="人力资源部" if role == "hr" else ["技术中心", "产品中心", "数据中心"][index % 3],
                is_active=True,
            )
        )
        missing -= 1
    db.commit()


def ensure_jobs(db) -> None:
    missing = max(0, TARGETS["jobs"] - count_rows(db, Job))
    if not missing:
        return
    owners = db.scalars(select(User).where(User.role.in_(["admin", "hr"])).order_by(User.id)).all()
    current_open = int(db.scalar(select(func.count()).select_from(Job).where(Job.status == "open")) or 0)
    specs = [
        ("数据分析师", "数据中心", "Python、SQL、数据可视化"),
        ("产品经理", "产品中心", "需求分析、招聘产品、跨团队协作"),
        ("招聘运营专员", "人力资源部", "招聘流程、企业微信、数据分析"),
        ("机器学习工程师", "AI 平台部", "Python、机器学习、模型部署"),
        ("测试开发工程师", "质量中心", "Python、自动化测试、接口测试"),
        ("DevOps 工程师", "技术中心", "Docker、Kubernetes、CI/CD"),
        ("数据产品经理", "数据中心", "数据产品、指标体系、SQL"),
        ("前端架构师", "技术中心", "TypeScript、React、工程化"),
        ("HRBP", "人力资源部", "组织发展、人才盘点、业务支持"),
        ("算法测试工程师", "质量中心", "模型评测、Python、数据质量"),
    ]
    now = utcnow()
    for index, (name, department, requirements) in enumerate(specs, start=1):
        if missing <= 0:
            break
        job_id = sid("job", index)
        if db.get(Job, job_id):
            continue
        is_open = current_open < 7
        opened = now - timedelta(days=90 - index * 4)
        db.add(
            Job(
                id=job_id,
                job_code=f"DEMO-2026-{index:03d}",
                job_name=name,
                department=department,
                description=f"虚构演示岗位：负责{name}相关工作与招聘数字化协作。",
                requirements=requirements,
                location=["上海", "杭州", "深圳", "北京"][index % 4],
                salary_min=Decimal(14_000 + index * 500),
                salary_max=Decimal(25_000 + index * 800),
                headcount=1 + index % 3,
                owner_id=owners[index % len(owners)].id,
                status="open" if is_open else "closed",
                opened_at=opened,
                closed_at=None if is_open else opened + timedelta(days=45),
            )
        )
        if is_open:
            current_open += 1
        missing -= 1
    db.commit()


def ensure_candidates(db) -> None:
    missing = max(0, TARGETS["candidates"] - count_rows(db, Candidate))
    if not missing:
        return
    creators = db.scalars(select(User).where(User.role == "hr").order_by(User.id)).all()
    sources = ["resume", "resume", "resume", "resume", "referral", "referral", "job_board", "job_board", "manual", "job_board"]
    now = utcnow()
    for index in range(1, 500):
        if missing <= 0:
            break
        candidate_id = sid("candidate", index)
        if db.get(Candidate, candidate_id):
            continue
        name = SURNAMES[(index - 1) % len(SURNAMES)] + GIVEN_NAMES[(index * 7) % len(GIVEN_NAMES)]
        skills = SKILL_GROUPS[index % len(SKILL_GROUPS)]
        created_at = now - timedelta(days=120 - index % 110, hours=index % 8)
        db.add(
            Candidate(
                id=candidate_id,
                name=name,
                phone=str(13910000000 + index),
                email=f"demo-candidate-{index:04d}@example.com",
                city=["上海", "杭州", "深圳", "北京", "苏州"][index % 5],
                source=sources[index % len(sources)],
                current_company=f"虚构科技{index % 18 + 1}有限公司",
                years_of_experience=Decimal(f"{1 + index % 12}.{index % 10}"),
                skills=skills,
                education=[{"school": f"虚构大学{index % 12 + 1}", "degree": "本科", "major": "计算机相关专业"}],
                work_experience=[{"company": f"虚构科技{index % 18 + 1}有限公司", "title": "专业岗位", "years": 1 + index % 6}],
                projects=[{"name": "招聘数字化演示项目", "description": "所有内容均为虚构数据"}],
                summary=f"{name}的虚构演示候选人档案，具备{'、'.join(skills)}相关经验。",
                created_by=creators[index % len(creators)].id,
                created_at=created_at,
                updated_at=created_at,
            )
        )
        missing -= 1
    db.commit()


def build_status_pool(db, needed: int) -> list[str]:
    current = Counter(dict(db.execute(select(CandidateJob.status, func.count()).group_by(CandidateJob.status)).all()))
    pool: list[str] = []
    for stage, target in STAGE_TARGETS.items():
        pool.extend([stage] * max(0, target - current.get(stage, 0)))
    while len(pool) < needed:
        pool.append(["new", "screening", "interview_1", "interview_2"][len(pool) % 4])
    RANDOM.shuffle(pool)
    return pool[:needed]


def ensure_applications(db) -> list[str]:
    missing = max(0, TARGETS["candidate_jobs"] - count_rows(db, CandidateJob))
    if not missing:
        return [sid("application", index) for index in range(1, 500) if db.get(CandidateJob, sid("application", index))]
    candidates = db.scalars(select(Candidate).where(Candidate.deleted_at.is_(None)).order_by(Candidate.created_at, Candidate.id)).all()
    open_jobs = db.scalars(select(Job).where(Job.status == "open").order_by(Job.id)).all()
    owners = db.scalars(select(User).where(User.role == "hr").order_by(User.id)).all()
    status_pool = build_status_pool(db, missing)
    existing_pairs = set(db.execute(select(CandidateJob.candidate_id, CandidateJob.job_id)).all())
    sources = ["resume", "resume", "referral", "job_board", "resume", "job_board", "manual"]
    now = utcnow()
    created_ids: list[str] = []
    pool_index = 0
    serial = 1
    for pass_number in (0, 1):
        for candidate_index, candidate in enumerate(candidates):
            if missing <= 0:
                break
            job = open_jobs[(candidate_index + pass_number * 3) % len(open_jobs)]
            if (candidate.id, job.id) in existing_pairs:
                continue
            app_id = sid("application", serial)
            while db.get(CandidateJob, app_id):
                serial += 1
                app_id = sid("application", serial)
            stage = status_pool[pool_index]
            chain = stage_chain(stage, serial)
            applied_at = now - timedelta(days=100 - serial % 92, hours=serial % 9)
            stage_at = applied_at + timedelta(days=max(0, len(chain) - 1) * 4)
            action = next_action(stage)
            db.add(
                CandidateJob(
                    id=app_id,
                    candidate_id=candidate.id,
                    job_id=job.id,
                    owner_id=owners[serial % len(owners)].id,
                    source=sources[serial % len(sources)],
                    status=stage,
                    match_score=Decimal(62 + serial % 36),
                    applied_at=applied_at,
                    stage_entered_at=stage_at,
                    next_action=action,
                    next_action_at=stage_at + timedelta(days=3 + serial % 5) if action else None,
                    rejection_reason="岗位匹配度未达到当前要求" if stage == "rejected" else None,
                    created_at=applied_at,
                    updated_at=stage_at,
                    version=len(chain),
                )
            )
            existing_pairs.add((candidate.id, job.id))
            created_ids.append(app_id)
            missing -= 1
            pool_index += 1
            serial += 1
        if missing <= 0:
            break
    db.commit()
    return created_ids


def generated_applications(db) -> list[CandidateJob]:
    ids = [sid("application", index) for index in range(1, 600)]
    return db.scalars(select(CandidateJob).where(CandidateJob.id.in_(ids)).order_by(CandidateJob.applied_at, CandidateJob.id)).all()


def ensure_approvals(db, apps: list[CandidateJob]) -> None:
    missing = max(0, TARGETS["approval_requests"] - count_rows(db, ApprovalRequest))
    if not missing:
        return
    current = Counter(dict(db.execute(select(ApprovalRequest.status, func.count()).group_by(ApprovalRequest.status)).all()))
    desired = {"pending": 10, "approved": 20, "rejected": 5}
    statuses: list[str] = []
    for status, target in desired.items():
        statuses.extend([status] * max(0, target - current.get(status, 0)))
    while len(statuses) < missing:
        statuses.append("pending")
    terminal = [a for a in apps if a.status in {"offer", "hired", "rejected", "withdrawn"}]
    nonterminal = [a for a in apps if a.status not in {"hired", "rejected", "withdrawn"}]
    decision_users = db.scalars(select(User).where(User.role.in_(["admin", "hr"])).order_by(User.id)).all()
    now = utcnow()
    for index, status in enumerate(statuses[:missing], start=1):
        approval_id = sid("approval", index)
        if db.get(ApprovalRequest, approval_id):
            continue
        app = terminal[index % len(terminal)] if status == "approved" else nonterminal[index % len(nonterminal)]
        target_status = app.status if status == "approved" else ("offer" if app.status == "final_interview" else "rejected")
        created_at = app.stage_entered_at + timedelta(hours=2)
        db.add(
            ApprovalRequest(
                id=approval_id,
                candidate_job_id=app.id,
                request_type="status_change",
                proposed_action="update_candidate_status",
                proposed_data={"target_status": target_status, "reason": "演示数据审批"},
                reason="高风险、低置信度或非常规阶段变更",
                confidence=Decimal("0.780") if status != "approved" else Decimal("0.930"),
                target_version=max(1, app.version - 1) if status == "approved" else app.version,
                idempotency_key=digest(f"demo-approval:{index}"),
                status=status,
                requested_by_type="agent",
                requested_by=decision_users[index % len(decision_users)].id,
                decided_by=decision_users[(index + 1) % len(decision_users)].id if status != "pending" else None,
                decision_comment="演示审批已核验" if status == "approved" else "演示审批未通过" if status == "rejected" else None,
                decided_at=created_at + timedelta(hours=4) if status != "pending" else None,
                created_at=created_at,
                updated_at=created_at + timedelta(hours=4) if status != "pending" else created_at,
            )
        )
    db.commit()


def ensure_inbound_events(db, apps: list[CandidateJob]) -> None:
    missing = max(0, TARGETS["inbound_events"] - count_rows(db, InboundEvent))
    if not missing:
        return
    pending_approvals = db.scalars(select(ApprovalRequest).where(ApprovalRequest.status == "pending").order_by(ApprovalRequest.id)).all()
    existing = set(db.scalars(select(InboundEvent.external_event_id).where(InboundEvent.source == "demo_seed")).all())
    created = 0
    for index in range(1, 1000):
        if created >= missing:
            break
        external_id = f"demo-seed-event-{index:04d}"
        if external_id in existing:
            continue
        app = apps[index % len(apps)]
        if index % 17 == 0:
            status, relevance = "failed", "relevant"
        elif index % 10 == 0 and pending_approvals:
            status, relevance = "pending_approval", "relevant"
        elif index % 7 == 0:
            status, relevance = "irrelevant", "irrelevant"
        else:
            status, relevance = "completed", "relevant"
        occurred_at = app.applied_at + timedelta(days=1 + index % 30, minutes=index % 55)
        approval = pending_approvals[index % len(pending_approvals)] if status == "pending_approval" else None
        db.add(
            InboundEvent(
                source="demo_seed",
                external_event_id=external_id,
                sender_external_id=f"demo_hr_{index % 5 + 1}",
                room_external_id=f"demo_recruitment_group_{index % 3 + 1}",
                message_type="text",
                content=(f"演示消息：候选人应聘记录 {app.id[:8]} 更新至 {app.status}" if relevance == "relevant" else "演示消息：下午茶已送达"),
                occurred_at=occurred_at,
                raw_payload={"demo_seed": True, "sequence": index},
                extracted_data={"candidate_job_id": app.id, "target_status": app.status} if relevance == "relevant" else None,
                relevance=relevance,
                confidence=Decimal("0.920") if status == "completed" else Decimal("0.620") if status == "pending_approval" else None,
                processing_status=status,
                approval_request_id=approval.id if approval else None,
                error_code="DEMO_PARSE_ERROR" if status == "failed" else None,
                error_message="演示解析异常" if status == "failed" else None,
                retry_count=3 if status == "failed" else 0,
                processed_at=occurred_at + timedelta(minutes=2) if status in {"completed", "irrelevant", "failed"} else None,
                created_at=occurred_at,
                updated_at=occurred_at + timedelta(minutes=2),
            )
        )
        created += 1
    db.commit()


def ensure_histories(db, apps: list[CandidateJob]) -> None:
    hr = db.scalar(select(User).where(User.role == "hr").order_by(User.id))
    approved = {
        item.candidate_job_id: item
        for item in db.scalars(select(ApprovalRequest).where(ApprovalRequest.status == "approved")).all()
        if item.candidate_job_id
    }
    events_by_app: dict[str, InboundEvent] = {}
    for event in db.scalars(select(InboundEvent).where(InboundEvent.processing_status == "completed").order_by(InboundEvent.id)).all():
        app_id = (event.extracted_data or {}).get("candidate_job_id")
        if app_id and app_id not in events_by_app:
            events_by_app[app_id] = event
    for app_index, app in enumerate(apps, start=1):
        exists = db.scalar(select(func.count()).select_from(CandidateStatusHistory).where(CandidateStatusHistory.candidate_job_id == app.id))
        if exists:
            continue
        chain = stage_chain(app.status, app_index)
        previous = None
        for step, stage in enumerate(chain):
            approval = approved.get(app.id) if step == len(chain) - 1 and stage in {"offer", "hired", "rejected", "withdrawn"} else None
            event = events_by_app.get(app.id) if step == len(chain) - 1 else None
            db.add(
                CandidateStatusHistory(
                    candidate_job_id=app.id,
                    from_status=previous,
                    to_status=stage,
                    reason="创建应聘关系" if previous is None else f"演示流程推进：{previous} → {stage}",
                    confidence=Decimal("0.930") if previous is not None else None,
                    operator_type="human" if previous is None else "agent",
                    operator_id=hr.id if previous is None else None,
                    approval_request_id=approval.id if approval else None,
                    inbound_event_id=event.id if event else None,
                    request_id=sid("history-request", f"{app.id}:{step}"),
                    created_at=app.applied_at + timedelta(days=step * 4),
                )
            )
            previous = stage
    db.commit()


def ensure_terminal_approvals(db) -> None:
    """Every high-risk terminal stage must be backed by one approved request."""
    decision_users = db.scalars(
        select(User).where(User.role.in_(["admin", "hr"])).order_by(User.id)
    ).all()
    high_risk_apps = db.scalars(
        select(CandidateJob)
        .where(CandidateJob.status.in_(["offer", "hired", "rejected", "withdrawn"]))
        .order_by(CandidateJob.id)
    ).all()
    for index, app in enumerate(high_risk_apps, start=1):
        existing = db.scalar(
            select(ApprovalRequest).where(
                ApprovalRequest.candidate_job_id == app.id,
                ApprovalRequest.status == "approved",
                ApprovalRequest.proposed_action == "update_candidate_status",
            )
        )
        history = db.scalar(
            select(CandidateStatusHistory)
            .where(
                CandidateStatusHistory.candidate_job_id == app.id,
                CandidateStatusHistory.to_status == app.status,
            )
            .order_by(CandidateStatusHistory.created_at.desc(), CandidateStatusHistory.id.desc())
        )
        if not history:
            raise RuntimeError(f"High-risk application {app.id} has no terminal history")
        if existing and (existing.proposed_data or {}).get("target_status") == app.status:
            if history.approval_request_id is None:
                history.approval_request_id = existing.id
            continue
        approval_id = sid("terminal-approval", app.id)
        approval = db.get(ApprovalRequest, approval_id)
        if not approval:
            approval = ApprovalRequest(
                id=approval_id,
                candidate_job_id=app.id,
                request_type="status_change",
                proposed_action="update_candidate_status",
                proposed_data={"target_status": app.status, "reason": "高风险阶段演示审批"},
                reason="Offer、入职、淘汰或退出必须经过人工审批",
                confidence=Decimal("0.930"),
                target_version=max(1, app.version - 1),
                idempotency_key=digest(f"demo-terminal-approval:{app.id}"),
                status="approved",
                requested_by_type="agent",
                requested_by=decision_users[index % len(decision_users)].id,
                decided_by=decision_users[(index + 1) % len(decision_users)].id,
                decision_comment="演示审批已核验并批准",
                decided_at=history.created_at,
                created_at=history.created_at - timedelta(hours=2),
                updated_at=history.created_at,
            )
            db.add(approval)
            db.flush()
        if history.approval_request_id is None:
            history.approval_request_id = approval.id
    db.commit()


def make_resume_bytes(candidate: Candidate, sequence: int) -> bytes:
    document = Document()
    document.add_heading(f"{candidate.name}的虚构演示简历", 0)
    document.add_paragraph(
        f"姓名：{candidate.name}\n手机号：{candidate.phone}\n邮箱：{candidate.email}\n"
        f"现居：{candidate.city}\n工作经验：{candidate.years_of_experience}年\n"
        f"技能：{'、'.join(candidate.skills)}\n"
        "本简历完全由 RecruitFlow 演示数据生成器创建，不对应任何真实个人。\n"
        f"演示序号：{sequence:04d}"
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def ensure_resumes(db) -> None:
    missing = max(0, TARGETS["resumes"] - count_rows(db, Resume))
    if not missing:
        return
    settings.ensure_directories()
    demo_dir = settings.upload_dir / "demo_seed"
    demo_dir.mkdir(parents=True, exist_ok=True)
    candidates = db.scalars(select(Candidate).where(Candidate.deleted_at.is_(None)).order_by(Candidate.created_at, Candidate.id)).all()
    uploaded_by = db.scalar(select(User).where(User.role == "hr").order_by(User.id))
    candidates_with_resume = set(db.scalars(select(Resume.candidate_id).where(Resume.candidate_id.is_not(None))).all())
    created = 0
    for index, candidate in enumerate(candidates, start=1):
        if created >= missing:
            break
        if candidate.id in candidates_with_resume:
            continue
        resume_id = sid("resume", candidate.id)
        if db.get(Resume, resume_id):
            continue
        content = make_resume_bytes(candidate, index)
        file_hash = hashlib.sha256(content).hexdigest()
        target = demo_dir / f"{resume_id}.docx"
        if not target.exists():
            target.write_bytes(content)
        parsed = {
            "name": candidate.name,
            "phone": candidate.phone,
            "email": candidate.email,
            "city": candidate.city,
            "current_company": candidate.current_company,
            "years_of_experience": str(candidate.years_of_experience),
            "skills": candidate.skills,
            "education": candidate.education,
            "work_experience": candidate.work_experience,
            "projects": candidate.projects,
            "summary": candidate.summary,
            "confidence": "0.950",
        }
        db.add(
            Resume(
                id=resume_id,
                candidate_id=candidate.id,
                uploaded_by=uploaded_by.id,
                original_file_name=f"{candidate.name}-虚构演示简历.docx",
                storage_path=str(target.resolve()),
                file_type="docx",
                file_sha256=file_hash,
                file_size=len(content),
                raw_text=f"{candidate.name}，{'、'.join(candidate.skills)}，虚构演示简历。",
                parsed_data=parsed,
                corrected_data=parsed,
                parse_status="completed",
                parser_model="demo_seed",
                created_at=candidate.created_at,
                updated_at=candidate.created_at + timedelta(minutes=2),
            )
        )
        candidates_with_resume.add(candidate.id)
        created += 1
    db.commit()


def interview_rounds_for(app: CandidateJob) -> list[str]:
    if app.status == "new":
        return []
    rank = {"screening": 0, "interview_1": 1, "interview_2": 2, "final_interview": 3, "on_hold": 0, "offer": 3, "hired": 3, "rejected": 2, "withdrawn": 1}.get(app.status, 0)
    return ["screening", "first", "second", "final"][: rank + 1]


def ensure_interviews(db, apps: list[CandidateJob]) -> None:
    missing = max(0, TARGETS["interviews"] - count_rows(db, Interview))
    if not missing:
        return
    interviewers = db.scalars(select(User).where(User.role == "interviewer").order_by(User.id)).all()
    creators = db.scalars(select(User).where(User.role == "hr").order_by(User.id)).all()
    now = utcnow().replace(minute=0, second=0, microsecond=0)
    created = 0
    slot = 0
    for app in apps:
        for round_name in interview_rounds_for(app):
            if created >= missing:
                break
            interview_id = sid("interview", f"{app.id}:{round_name}")
            if db.get(Interview, interview_id):
                continue
            is_completed = slot % 4 == 0
            scheduled_at = now - timedelta(days=1 + slot % 20) if is_completed else now + timedelta(days=1 + slot % 14, hours=slot % 7)
            recommendation = ["pass", "hold", "reject"][slot % 3] if is_completed else "pending"
            db.add(
                Interview(
                    id=interview_id,
                    candidate_job_id=app.id,
                    round=round_name,
                    interview_type=["online", "onsite", "phone"][slot % 3],
                    scheduled_at=scheduled_at,
                    duration_minutes=[45, 60, 75][slot % 3],
                    interviewer_id=interviewers[slot % len(interviewers)].id,
                    additional_interviewers=[],
                    meeting_url=f"https://meeting.example.com/demo/{interview_id[:8]}" if slot % 3 == 0 else None,
                    location="上海演示会议室" if slot % 3 == 1 else None,
                    status="completed" if is_completed else "scheduled",
                    strengths="专业基础扎实，沟通清晰" if is_completed else None,
                    weaknesses="部分场景经验需要进一步确认" if is_completed else None,
                    feedback="虚构面试反馈，仅用于演示" if is_completed else None,
                    ai_summary="候选人具备岗位基础能力，建议结合后续面试综合判断。" if is_completed else None,
                    recommendation=recommendation,
                    feedback_submitted_at=scheduled_at + timedelta(hours=1) if is_completed else None,
                    created_by=creators[slot % len(creators)].id,
                    created_at=min(app.applied_at + timedelta(days=1), scheduled_at - timedelta(days=1)),
                    updated_at=scheduled_at + timedelta(hours=1) if is_completed else scheduled_at - timedelta(days=1),
                )
            )
            created += 1
            slot += 1
        if created >= missing:
            break
    db.commit()


def repair_invalid_interviews(db) -> None:
    stage_by_round = {
        "screening": "screening",
        "first": "interview_1",
        "second": "interview_2",
        "final": "final_interview",
    }
    interviews = db.scalars(select(Interview).order_by(Interview.id)).all()
    for interview in interviews:
        required_stage = stage_by_round[interview.round]
        has_stage = db.scalar(
            select(func.count())
            .select_from(CandidateStatusHistory)
            .where(
                CandidateStatusHistory.candidate_job_id == interview.candidate_job_id,
                CandidateStatusHistory.to_status == required_stage,
            )
        )
        if has_stage:
            continue
        target = db.scalar(
            select(CandidateJob)
            .where(
                CandidateJob.id.in_(
                    select(CandidateStatusHistory.candidate_job_id).where(
                        CandidateStatusHistory.to_status == required_stage
                    )
                ),
                ~CandidateJob.id.in_(
                    select(Interview.candidate_job_id).where(Interview.round == interview.round)
                ),
            )
            .order_by(CandidateJob.id)
        )
        if not target:
            raise RuntimeError(f"No valid application available for interview {interview.id}")
        interview.candidate_job_id = target.id
        for notification in db.scalars(
            select(Notification).where(Notification.interview_id == interview.id)
        ).all():
            notification.candidate_job_id = target.id
    db.commit()


def ensure_notifications(db, apps: list[CandidateJob]) -> None:
    missing = max(0, TARGETS["notifications"] - count_rows(db, Notification))
    if not missing:
        return
    interviews = db.scalars(select(Interview).order_by(Interview.scheduled_at, Interview.id)).all()
    created = 0
    now = utcnow()
    for index in range(1, 1000):
        if created >= missing:
            break
        notification_id = sid("notification", index)
        if db.get(Notification, notification_id):
            continue
        interview = interviews[index % len(interviews)]
        app = db.get(CandidateJob, interview.candidate_job_id)
        scheduled_at = interview.scheduled_at - timedelta(hours=24)
        sent = scheduled_at < now
        db.add(
            Notification(
                id=notification_id,
                candidate_job_id=app.id,
                interview_id=interview.id,
                channel="in_app" if index % 4 else "wecom_bot",
                recipient_type="user",
                recipient=interview.interviewer_id,
                content=f"演示提醒：{interview.round}面试将于{interview.scheduled_at:%m月%d日 %H:%M}进行",
                scheduled_at=scheduled_at,
                sent_at=scheduled_at + timedelta(minutes=1) if sent else None,
                status="sent" if sent else "pending",
                retry_count=0,
                created_at=scheduled_at - timedelta(days=1),
                updated_at=scheduled_at + timedelta(minutes=1) if sent else scheduled_at - timedelta(days=1),
            )
        )
        created += 1
    db.commit()


def ensure_communications(db, apps: list[CandidateJob]) -> None:
    missing = max(0, TARGETS["communications"] - count_rows(db, Communication))
    if not missing:
        return
    operators = db.scalars(select(User).where(User.role == "hr").order_by(User.id)).all()
    events = db.scalars(select(InboundEvent).where(InboundEvent.source == "demo_seed").order_by(InboundEvent.id)).all()
    existing_sequences = set()
    for content in db.scalars(select(Communication.content).where(Communication.content.like("演示数据:%"))).all():
        if content:
            existing_sequences.add(content.split(":", 2)[1])
    created = 0
    channels = ["wecom", "phone", "email", "interview"]
    intents = ["status_update", "interview_schedule", "follow_up", "candidate_question"]
    for index in range(1, 2000):
        if created >= missing:
            break
        sequence = f"{index:04d}"
        if sequence in existing_sequences:
            continue
        app = apps[index % len(apps)]
        event = events[index % len(events)] if events else None
        channel = channels[index % len(channels)]
        occurred_at = app.applied_at + timedelta(days=1 + index % 45, minutes=index % 50)
        db.add(
            Communication(
                candidate_id=app.candidate_id,
                candidate_job_id=app.id,
                inbound_event_id=event.id if event and channel == "wecom" else None,
                channel=channel,
                direction="inbound" if index % 3 == 0 else "outbound",
                content=f"演示数据:{sequence}:候选人沟通记录，仅用于系统功能展示",
                summary=f"{channel}渠道的虚构招聘沟通",
                intent=intents[index % len(intents)],
                next_action=next_action(app.status),
                operator_id=operators[index % len(operators)].id,
                occurred_at=occurred_at,
                created_at=occurred_at,
            )
        )
        created += 1
    db.commit()


def ensure_agent_data(db) -> None:
    users = db.scalars(select(User).where(User.role.in_(["admin", "hr"])).order_by(User.id)).all()
    missing_conversations = max(0, TARGETS["agent_conversations"] - count_rows(db, AgentConversation))
    for index in range(1, 200):
        if missing_conversations <= 0:
            break
        conversation_id = sid("conversation", index)
        if db.get(AgentConversation, conversation_id):
            continue
        created_at = utcnow() - timedelta(days=30 - index % 28, hours=index % 9)
        db.add(
            AgentConversation(
                id=conversation_id,
                user_id=users[index % len(users)].id,
                title=["本周候选人跟进", "岗位招聘漏斗分析", "面试安排检查", "待审批事项", "候选人搜索"][index % 5],
                status="active",
                created_at=created_at,
                updated_at=created_at + timedelta(minutes=20),
            )
        )
        missing_conversations -= 1
    db.commit()

    missing_messages = max(0, TARGETS["agent_messages"] - count_rows(db, AgentMessage))
    conversations = db.scalars(select(AgentConversation).order_by(AgentConversation.created_at, AgentConversation.id)).all()
    message_pairs = [
        ("本周有哪些候选人需要跟进？", "当前共有 18 位候选人需要跟进，已按紧急程度排序。", "search_candidates"),
        ("查看当前招聘漏斗", "已汇总开放岗位、各阶段人数和待审批事项。", "get_recruitment_dashboard"),
        ("查询待审批记录", "当前高风险操作均已进入人工审批队列。", "get_recruitment_dashboard"),
    ]
    conversation_index = 0
    while missing_messages > 0:
        conversation = conversations[conversation_index % len(conversations)]
        existing_count = int(db.scalar(select(func.count()).select_from(AgentMessage).where(AgentMessage.conversation_id == conversation.id)) or 0)
        pair = message_pairs[(existing_count // 2) % len(message_pairs)]
        role = "user" if existing_count % 2 == 0 else "assistant"
        content = pair[0] if role == "user" else pair[1]
        db.add(
            AgentMessage(
                conversation_id=conversation.id,
                role=role,
                content=content,
                tool_name=pair[2] if role == "assistant" else None,
                tool_input={"demo_seed": True} if role == "assistant" else None,
                tool_output={"success": True, "demo_seed": True} if role == "assistant" else None,
                status="completed",
                request_id=sid("agent-message-request", f"{conversation.id}:{existing_count}"),
                created_at=conversation.created_at + timedelta(minutes=existing_count * 2),
            )
        )
        missing_messages -= 1
        conversation_index += 1
        if conversation_index % len(conversations) == 0:
            db.flush()
    db.commit()


def ensure_sync_jobs(db, apps: list[CandidateJob]) -> None:
    missing = max(0, TARGETS["document_sync_jobs"] - count_rows(db, DocumentSyncJob))
    if not missing:
        return
    existing_keys = set(db.scalars(select(DocumentSyncJob.idempotency_key)).all())
    created = 0
    now = utcnow()
    for index in range(1, 2000):
        if created >= missing:
            break
        key = digest(f"demo-sync:{index}")
        if key in existing_keys:
            continue
        app = apps[index % len(apps)]
        if index % 50 == 0:
            status = "failed"
        elif index % 20 == 0:
            status = "pending"
        else:
            status = "succeeded"
        db.add(
            DocumentSyncJob(
                sink_type="local_excel",
                entity_type="application",
                entity_id=app.id,
                operation="upsert",
                payload={
                    "candidate_id": app.candidate_id,
                    "job_id": app.job_id,
                    "status": app.status,
                    "owner_id": app.owner_id,
                    "next_action": app.next_action,
                    "updated_at": app.updated_at.isoformat(),
                },
                payload_version=1,
                idempotency_key=key,
                status=status,
                retry_count=3 if status == "failed" else 0,
                external_record_id=f"演示招聘数据!{index + 1}" if status == "succeeded" else None,
                error_code="DEMO_SYNC_FAILURE" if status == "failed" else None,
                error_message="演示同步失败，可在界面中查看重试状态" if status == "failed" else None,
                last_attempt_at=now - timedelta(minutes=index % 30) if status != "pending" else None,
                completed_at=now - timedelta(minutes=index % 30) if status == "succeeded" else None,
                created_at=now - timedelta(days=index % 60),
                updated_at=now - timedelta(minutes=index % 30),
            )
        )
        existing_keys.add(key)
        created += 1
    db.commit()


def ensure_audit_logs(db, apps: list[CandidateJob]) -> None:
    missing = max(0, TARGETS["audit_logs"] - count_rows(db, AuditLog))
    if not missing:
        return
    users = db.scalars(select(User).where(User.role.in_(["admin", "hr"])).order_by(User.id)).all()
    existing_requests = set(db.scalars(select(AuditLog.request_id).where(AuditLog.source == "demo_seed")).all())
    actions = ["candidate.create", "application.create", "application.status_change", "interview.create", "interview.feedback", "approval.create", "document.sync"]
    created = 0
    now = utcnow()
    for index in range(1, 5000):
        if created >= missing:
            break
        request = sid("audit-request", index)
        if request in existing_requests:
            continue
        app = apps[index % len(apps)]
        action = actions[index % len(actions)]
        db.add(
            AuditLog(
                request_id=request,
                user_id=users[index % len(users)].id if index % 4 else None,
                actor_type="agent" if index % 4 == 0 else "human" if index % 3 else "system",
                action=action,
                target_type="candidate_job",
                target_id=app.id,
                before_data={"status": "screening"} if action == "application.status_change" else None,
                after_data={"status": app.status, "demo_seed": True},
                source="demo_seed",
                ip_address="127.0.0.1",
                created_at=now - timedelta(days=index % 90, minutes=index % 55),
            )
        )
        existing_requests.add(request)
        created += 1
    db.commit()


def audit_consistency() -> dict[str, int]:
    results: dict[str, int] = {}
    with engine.connect() as connection:
        foreign_keys = connection.execute(
            text(
                """
                SELECT table_name,column_name,referenced_table_name,referenced_column_name
                FROM information_schema.key_column_usage
                WHERE table_schema=:db AND referenced_table_name IS NOT NULL
                """
            ),
            {"db": settings.db_name},
        ).all()
        orphan_total = 0
        for table, column, parent, parent_column in foreign_keys:
            orphan_total += int(
                connection.execute(
                    text(
                        f"SELECT COUNT(*) FROM `{table}` ch LEFT JOIN `{parent}` p "
                        f"ON ch.`{column}`=p.`{parent_column}` "
                        f"WHERE ch.`{column}` IS NOT NULL AND p.`{parent_column}` IS NULL"
                    )
                ).scalar_one()
            )
        results["foreign_key_orphans"] = orphan_total
        checks = {
            "applications_without_history": "SELECT COUNT(*) FROM candidate_jobs cj WHERE NOT EXISTS (SELECT 1 FROM candidate_status_history h WHERE h.candidate_job_id=cj.id)",
            "latest_history_mismatch": "SELECT COUNT(*) FROM candidate_jobs cj JOIN candidate_status_history h ON h.id=(SELECT h2.id FROM candidate_status_history h2 WHERE h2.candidate_job_id=cj.id ORDER BY h2.created_at DESC,h2.id DESC LIMIT 1) WHERE cj.status<>h.to_status",
            "duplicate_candidate_job": "SELECT COUNT(*) FROM (SELECT candidate_id,job_id,COUNT(*) n FROM candidate_jobs GROUP BY candidate_id,job_id HAVING n>1) x",
            "duplicate_resume_hash": "SELECT COUNT(*) FROM (SELECT file_sha256,COUNT(*) n FROM resumes GROUP BY file_sha256 HAVING n>1) x",
            "duplicate_inbound_event": "SELECT COUNT(*) FROM (SELECT source,external_event_id,COUNT(*) n FROM inbound_events GROUP BY source,external_event_id HAVING n>1) x",
            "duplicate_sync_key": "SELECT COUNT(*) FROM (SELECT idempotency_key,COUNT(*) n FROM document_sync_jobs GROUP BY idempotency_key HAVING n>1) x",
            "broken_history_chain": "SELECT COUNT(*) FROM candidate_status_history h JOIN candidate_status_history prev ON prev.id=(SELECT h2.id FROM candidate_status_history h2 WHERE h2.candidate_job_id=h.candidate_job_id AND (h2.created_at<h.created_at OR (h2.created_at=h.created_at AND h2.id<h.id)) ORDER BY h2.created_at DESC,h2.id DESC LIMIT 1) WHERE h.from_status<>prev.to_status",
            "high_risk_without_approval": "SELECT COUNT(*) FROM candidate_jobs cj WHERE cj.status IN ('offer','hired','rejected','withdrawn') AND NOT EXISTS (SELECT 1 FROM approval_requests a WHERE a.candidate_job_id=cj.id AND a.status='approved' AND JSON_UNQUOTE(JSON_EXTRACT(a.proposed_data,'$.target_status'))=cj.status)",
            "approved_target_mismatch": "SELECT COUNT(*) FROM approval_requests a JOIN candidate_jobs cj ON cj.id=a.candidate_job_id WHERE a.status='approved' AND a.proposed_action='update_candidate_status' AND JSON_UNQUOTE(JSON_EXTRACT(a.proposed_data,'$.target_status'))<>cj.status",
            "pending_target_already_applied": "SELECT COUNT(*) FROM approval_requests a JOIN candidate_jobs cj ON cj.id=a.candidate_job_id WHERE a.status='pending' AND a.proposed_action='update_candidate_status' AND JSON_UNQUOTE(JSON_EXTRACT(a.proposed_data,'$.target_status'))=cj.status",
            "approved_without_history_link": "SELECT COUNT(*) FROM approval_requests a WHERE a.status='approved' AND a.proposed_action='update_candidate_status' AND NOT EXISTS (SELECT 1 FROM candidate_status_history h WHERE h.approval_request_id=a.id)",
            "interview_without_stage": "SELECT COUNT(*) FROM interviews i WHERE NOT EXISTS (SELECT 1 FROM candidate_status_history h WHERE h.candidate_job_id=i.candidate_job_id AND h.to_status=CASE i.round WHEN 'screening' THEN 'screening' WHEN 'first' THEN 'interview_1' WHEN 'second' THEN 'interview_2' WHEN 'final' THEN 'final_interview' END)",
        }
        for name, sql in checks.items():
            results[name] = int(connection.execute(text(sql)).scalar_one())
    return results


def table_counts(db) -> dict[str, int]:
    models = [
        User, Job, Candidate, CandidateJob, Resume, Interview, CandidateStatusHistory,
        Communication, Notification, AuditLog, AgentConversation, AgentMessage,
        ApprovalRequest, InboundEvent, DocumentSyncJob,
    ]
    return {model.__tablename__: count_rows(db, model) for model in models}


def main() -> None:
    settings.validate_database_boundary()
    if settings.db_name != "hr_recruitment":
        raise RuntimeError("Demo data generation is restricted to hr_recruitment")
    RANDOM.seed(SEED)
    with SessionLocal() as db:
        before = table_counts(db)
        ensure_users(db)
        ensure_jobs(db)
        ensure_candidates(db)
        ensure_applications(db)
        apps = generated_applications(db)
        if not apps:
            apps = db.scalars(select(CandidateJob).order_by(CandidateJob.id)).all()
        ensure_approvals(db, apps)
        ensure_inbound_events(db, apps)
        ensure_histories(db, apps)
        ensure_terminal_approvals(db)
        ensure_resumes(db)
        ensure_interviews(db, apps)
        repair_invalid_interviews(db)
        ensure_notifications(db, apps)
        ensure_communications(db, apps)
        ensure_agent_data(db)
        ensure_sync_jobs(db, apps)
        ensure_audit_logs(db, apps)
        after = table_counts(db)
    audit = audit_consistency()
    print("DEMO_DATA_GENERATION_OK")
    for table, count in after.items():
        print(f"{table}: {before[table]} -> {count}")
    print("CONSISTENCY_AUDIT")
    for name, value in audit.items():
        print(f"{name}: {value}")
    if any(audit.values()):
        raise RuntimeError(f"Consistency audit failed: {audit}")


if __name__ == "__main__":
    main()
