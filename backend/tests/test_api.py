from __future__ import annotations

import uuid
from io import BytesIO

import pytest
from docx import Document
from sqlalchemy import inspect

from app.agent.graph import TOOL_WHITELIST, build_tools
from app.database import SessionLocal, engine
from app.models import User


def test_schema_has_exact_core_tables():
    tables = set(inspect(engine).get_table_names())
    expected = {
        "users", "jobs", "candidates", "candidate_jobs", "resumes", "interviews",
        "candidate_status_history", "communications", "notifications", "audit_logs",
        "agent_conversations", "agent_messages", "approval_requests", "inbound_events",
        "document_sync_jobs", "agent_user_preferences", "agent_tool_runs",
    }
    assert expected.issubset(tables)


def test_health_and_auth(client):
    assert client.get("/health/live").status_code == 200
    assert client.get("/health/ready").status_code == 200
    assert client.get("/api/dashboard").status_code == 401


def test_candidates_are_masked(hr_client):
    response = hr_client.get("/api/candidates")
    assert response.status_code == 200
    data = response.json()["data"]
    assert data["total"] >= 12
    assert "****" in data["items"][0]["phone"]


def test_inbound_message_is_idempotent(hr_client):
    event_id = "pytest-inbound-idempotent-001"
    body = {"external_event_id": event_id, "content": "李明一面通过，安排二面"}
    first = hr_client.post("/api/inbound/demo", json=body)
    second = hr_client.post("/api/inbound/demo", json=body)
    assert first.status_code == 200
    assert second.status_code == 200
    assert second.json()["data"]["duplicate"] is True
    assert first.json()["data"]["event_id"] == second.json()["data"]["event_id"]


def test_high_risk_message_creates_approval(hr_client):
    body = {"external_event_id": "pytest-high-risk-001", "content": "王芳终面不通过，建议淘汰"}
    response = hr_client.post("/api/inbound/demo", json=body)
    assert response.status_code == 200
    assert response.json()["data"]["status"] == "pending_approval"


def test_agent_uses_whitelisted_dashboard_tool(hr_client):
    response = hr_client.post(
        "/api/agent/chat",
        json={
            "message": "目前有多少开放岗位和待审批？",
            "idempotency_key": "pytest-agent-" + uuid.uuid4().hex,
        },
    )
    assert response.status_code == 200, response.text
    data = response.json()["data"]
    assert data["tool_calls"][0]["name"] == "get_recruitment_dashboard"
    assert data["status"] == "completed"


def test_sidebar_has_two_groups_and_conversation_can_be_renamed(hr_client):
    page = hr_client.get("/agent")
    assert page.status_code == 200
    assert page.text.count('class="nav-group-toggle') == 2
    assert "AI 招聘助手" in page.text
    assert "招聘信息" in page.text
    assert "招聘看板" in page.text

    created = hr_client.post(
        "/api/agent/chat",
        json={
            "message": "查询招聘看板作为会话重命名测试",
            "idempotency_key": "rename-test-" + uuid.uuid4().hex,
        },
    )
    assert created.status_code == 200, created.text
    conversation_id = created.json()["data"]["conversation_id"]
    title = "我的招聘数据会话"
    renamed = hr_client.patch(
        f"/api/agent/conversations/{conversation_id}", json={"title": title}
    )
    assert renamed.status_code == 200, renamed.text
    conversations = hr_client.get("/api/agent/conversations").json()["data"]
    assert conversations[0]["id"] == conversation_id
    assert conversations[0]["title"] == title


def test_visible_business_enums_are_translated_to_chinese(hr_client):
    candidates = hr_client.get("/candidates")
    interviews = hr_client.get("/interviews")
    approvals = hr_client.get("/approvals")
    assert candidates.status_code == interviews.status_code == approvals.status_code == 200
    assert ">二面<" in candidates.text
    assert ">内部推荐<" in candidates.text
    assert ">线上面试<" in interviews.text
    assert "更新候选人阶段" in approvals.text
    assert ">interview_2<" not in candidates.text
    assert ">job_board<" not in candidates.text
    assert ">online<" not in interviews.text


def test_agent_registers_the_complete_tool_whitelist():
    with SessionLocal() as db:
        user = db.query(User).filter(User.username == "hr_demo").one()
        tools = build_tools(db, user)
        assert set(tools) == set(TOOL_WHITELIST)
        parsed = tools["parse_resume"].invoke(
            {"resume_text": "姓名：工具测试人，3年工作经验，熟悉 Python 和 FastAPI"}
        )
        assert parsed["parser"] == "rule"
        assert "Python" in parsed["parsed"]["skills"]


def test_missing_csrf_is_rejected(client):
    clean = client.headers.pop("X-CSRF-Token", None)
    response = client.post(
        "/api/inbound/demo", json={"external_event_id": "csrf-test", "content": "李明二面"}
    )
    if clean:
        client.headers["X-CSRF-Token"] = clean
    assert response.status_code == 403


def test_docx_resume_parses_without_llm_key(hr_client):
    jobs = hr_client.get("/api/jobs?status=open").json()["data"]
    document = Document()
    document.add_heading("虚构候选人简历", 0)
    document.add_paragraph(
        "姓名：测试候选人\n手机号：13800009999\n邮箱：resume-test@example.com\n"
        "拥有 4 年工作经验，熟悉 Python、FastAPI、MySQL、Docker。"
        "曾负责招聘数据自动化系统，包含接口设计、状态机、消息处理和可视化看板。"
    )
    buffer = BytesIO()
    document.save(buffer)
    response = hr_client.post(
        "/api/candidates/import-resume",
        data={"job_id": jobs[0]["id"]},
        files={"file": (f"resume-{uuid.uuid4().hex}.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 200, response.text
    data = response.json()["data"]
    assert data["parse_status"] == "review_required"
    assert "Python" in data["parsed_data"]["skills"]


def test_png_resume_import_exposes_ocr_metadata(hr_client, monkeypatch):
    from PIL import Image
    from app.ocr import OcrPageResult

    buffer = BytesIO()
    Image.new("RGB", (240, 160), "white").save(buffer, format="PNG")

    class StubClient:
        def recognize_image(self, payload):
            return OcrPageResult(
                "姓名：图片测试人\n手机号：13900008888\n邮箱：image-test@example.com\n"
                "拥有 5 年工作经验，熟悉 Python、FastAPI、MySQL、Docker。" * 3,
                0.93,
            )

    monkeypatch.setattr("app.services.get_aliyun_ocr_client", lambda: StubClient())
    jobs = hr_client.get("/api/jobs?status=open").json()["data"]
    response = hr_client.post(
        "/api/candidates/import-resume",
        data={"job_id": jobs[0]["id"]},
        files={"file": (f"image-{uuid.uuid4().hex}.png", buffer.getvalue(), "image/png")},
    )
    assert response.status_code == 200, response.text
    data = response.json()["data"]
    assert data["extraction_method"] == "aliyun_ocr"
    assert data["page_count"] == 1
    assert data["ocr_confidence"] == pytest.approx(0.93)
