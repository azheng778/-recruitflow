from __future__ import annotations

import os
import uuid
from io import BytesIO

import pytest
from docx import Document
from sqlalchemy import text

from app.config import settings
from app.database import engine
from app.database import SessionLocal
from app.models import AgentToolRun


pytestmark = pytest.mark.skipif(
    os.getenv("RUN_REAL_LLM_TESTS") != "1",
    reason="真实 LLM 集成测试默认关闭；设置 RUN_REAL_LLM_TESTS=1 后单独运行",
)


def test_real_database_and_qwen_resume_api(hr_client):
    assert settings.llm_api_key, "LLM_API_KEY 未配置"
    assert settings.db_name == settings.test_db_name
    assert settings.db_name != "langchain_db"

    with engine.connect() as connection:
        assert connection.scalar(text("SELECT DATABASE()")) == settings.test_db_name
        assert connection.scalar(text("SELECT 1")) == 1

    unique = uuid.uuid4().hex[:10]
    document = Document()
    document.add_heading("真实模型集成测试简历", 0)
    document.add_paragraph(
        "姓名：林海。手机号：13900001234。邮箱：linhai.integration@example.com。"
        "现居上海，拥有5年软件工程经验，目前任职于虚构科技有限公司。"
        "熟悉 Python、FastAPI、LangChain、LangGraph、MySQL 和 Docker，"
        "负责过招聘数据自动化、候选人状态机、企业消息处理、审批工作流、"
        "数据可视化看板和大语言模型工具调用项目。"
        "教育经历为虚构大学计算机科学本科。所有内容均为自动化测试虚构数据。"
        f"测试批次：{unique}。"
    )
    buffer = BytesIO()
    document.save(buffer)

    jobs = hr_client.get("/api/jobs?status=open").json()["data"]
    response = hr_client.post(
        "/api/candidates/import-resume",
        data={"job_id": jobs[0]["id"]},
        files={
            "file": (
                f"real-llm-{unique}.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200, response.text
    result = response.json()["data"]
    parsed = result["parsed_data"]
    assert result["parse_status"] == "review_required"
    assert parsed["name"] == "林海"
    assert "Python" in parsed["skills"]
    assert float(parsed["confidence"]) >= 0
    print(
        {
            "database": settings.test_db_name,
            "model": settings.llm_model,
            "resume_id": result["resume_id"],
            "name": parsed["name"],
            "skills": parsed["skills"],
            "confidence": parsed["confidence"],
        }
    )


def test_real_qwen_agent_structured_intent_and_tool_call(hr_client):
    key = "real-agent-" + uuid.uuid4().hex
    response = hr_client.post(
        "/api/agent/chat",
        json={
            "message": "请给我看当前招聘漏斗，并告诉我开放岗位和待审批数量",
            "idempotency_key": key,
            "client_timezone": "Asia/Shanghai",
        },
    )
    assert response.status_code == 200, response.text
    result = response.json()["data"]
    assert result["status"] == "completed"
    assert result["intent"]["name"] == "dashboard_query"
    assert result["tool_calls"][0]["name"] == "get_recruitment_dashboard"
    with SessionLocal() as db:
        run = db.query(AgentToolRun).filter(
            AgentToolRun.conversation_id == result["conversation_id"]
        ).one()
        assert run.model_name == settings.llm_model
        assert run.prompt_version == settings.agent_prompt_version
